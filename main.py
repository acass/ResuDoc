import streamlit as st
import docx
from docx import Document
import io
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openai
import os
from io import BytesIO
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Set the page config and title
st.set_page_config(page_title="Resume Optimizer", layout="wide")
st.title("🚀 Resume Optimizer")

# Description of the app
st.markdown("""
This app helps you tailor your resume to a specific job description. 
Upload your current resume, enter the job description, and get a customized resume that highlights 
relevant skills and experiences.
""")

# Get OpenAI API Key from environment variables or allow user input as fallback
openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key:
    openai_api_key = st.sidebar.text_input("OpenAI API Key", type="password")
    if openai_api_key:
        os.environ["OPENAI_API_KEY"] = openai_api_key
else:
    st.sidebar.success("✅ OpenAI API Key loaded from environment variables")


def extract_header(doc):
    """Extract the header (contact information) from the resume."""
    # Assume header is in the first few paragraphs
    header_text = []
    for i, para in enumerate(doc.paragraphs):
        if i < 5 and para.text.strip():  # Adjust the number as needed
            header_text.append(para.text)
        elif i >= 5 and any(keyword in para.text.lower() for keyword in ["experience", "education", "skills"]):
            break
    return header_text


def extract_content(doc):
    """Extract the main content from the resume."""
    content = []
    header_length = len(extract_header(doc))

    for i, para in enumerate(doc.paragraphs):
        if i >= header_length and para.text.strip():
            content.append(para.text)

    return "\n".join(content)

def optimize_resume(resume_content, job_description):
    """Use OpenAI to optimize the resume based on the job description."""
    if not openai_api_key:
        st.error("Please enter your OpenAI API key in the sidebar or use a .env file")
        return None

    client = openai.OpenAI(api_key=openai_api_key)

    prompt = f"""
    You are a professional ATS-optimization expert and resume writer specializing in tailoring resumes to specific job descriptions. Your task is to rewrite the provided resume to maximize ATS scoring and match the job description while maintaining truthfulness.

    ## Resume Content:
    {resume_content}

    ## Job Description:
    {job_description}

    ## Your Task:
    Rewrite the resume content to better align with the job description by:

    1. Identifying key skills, qualifications, and technologies in the job description and highlighting matching elements in the resume
    2. Using the exact same keywords and phrases from the job description when they genuinely match the candidate's experience
    3. Reorganizing bullet points to prioritize the most relevant experiences first
    4. Quantifying achievements with metrics where possible (maintain any existing metrics)
    5. Adjusting the professional summary/objective to specifically target this position
    6. Using action verbs that demonstrate impact
    7. Removing or downplaying irrelevant experiences
    8. Maintaining the same overall sections (Education, Experience, Skills, etc.)
    9. Keeping the same factual information - do not fabricate experiences or skills

    ## Format Requirements:
    - Maintain professional, concise language
    - Use bullet points for experiences
    - Keep the same chronological order of positions
    - Preserve all contact information
    - Focus on achievements rather than responsibilities
    - Ensure each bullet point demonstrates a skill relevant to the job description

    Return ONLY the optimized resume content without explanations, commentary, or additional text.
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert resume writer specializing in ATS optimization."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.5,
            max_tokens=2000
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error optimizing resume: {e}")
        return None

def create_optimized_resume(original_doc, optimized_content):
    """Create a new Word document with the original header and optimized content."""
    new_doc = Document()

    # Copy the header from the original document
    header_paragraphs = extract_header(original_doc)
    for header_text in header_paragraphs:
        header_para = new_doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run(header_text)
        run.font.size = Pt(12)
        run.font.bold = True

    # Add a separator
    new_doc.add_paragraph()

    # Add the optimized content
    content_sections = optimized_content.split('\n')
    for section in content_sections:
        if section.strip():
            # Check if it's a section header
            if any(section.lower().startswith(keyword) for keyword in
                   ["experience", "education", "skills", "summary", "objective"]):
                para = new_doc.add_paragraph()
                run = para.add_run(section)
                run.font.bold = True
                run.font.size = Pt(14)
            else:
                para = new_doc.add_paragraph(section)
                para.style = 'Normal'

    return new_doc

# Create two columns for input fields
col1, col2 = st.columns(2)

with col1:
    # Resume file uploader
    st.header("📄 Upload Your Resume")
    resume_file = st.file_uploader("Upload your current resume (Word format)", type=['docx'])

with col2:
    # Job description input
    st.header("📝 Job Description")
    job_description = st.text_area("Enter the job description", height=300)

# Process button
if st.button("🔄 Optimize Resume"):
    if resume_file is None:
        st.error("Please upload your resume file")
    elif not job_description:
        st.error("Please enter the job description")
    elif not openai_api_key:
        st.error("Please enter your OpenAI API key in the sidebar or use a .env file")
    else:
        with st.spinner("Processing..."):
            # Read the resume file
            doc = Document(resume_file)

            # Extract content
            resume_content = extract_content(doc)

            # Optimize the resume
            optimized_content = optimize_resume(resume_content, job_description)

            if optimized_content:
                # Create new document
                new_doc = create_optimized_resume(doc, optimized_content)

                # Save the document
                doc_io = BytesIO()
                new_doc.save(doc_io)
                doc_io.seek(0)

                # Display download button with a non-empty label
                st.success("✅ Resume optimized successfully!")
                st.download_button(
                    label="⬇️ Download Optimized Resume",
                    data=doc_io,
                    file_name="optimized_resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                # Preview of changes
                st.subheader("Preview of Optimized Resume")
                st.text_area("Preview", optimized_content, height=400)

# Footer
st.markdown("---")
st.markdown("Resume Optimizer by Allen Cass - Powered by Streamlit and OpenAI")



